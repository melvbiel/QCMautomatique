import QCM
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def format_paysage(doc):
    section = doc.sections[0]
    section.orientation = 1  # Paysage
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

def set_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')  # Deux colonnes
    cols.set(qn('w:space'), '720')  # Espacement entre les colonnes
    sectPr.append(cols)

def set_line_spacing(paragraph, spacing=1.0, before=0, after=0):
    p_pr = paragraph._element.get_or_add_pPr()
    spacing_elem = OxmlElement('w:spacing')
    spacing_elem.set(qn('w:line'), str(int(spacing * 240)))  # 1.0 interligne
    spacing_elem.set(qn('w:before'), str(before))  # Réduction espace avant
    spacing_elem.set(qn('w:after'), str(after))  # Réduction espace après
    p_pr.append(spacing_elem)

def create_qcm():
    questions = QCM.lire_questionnaire('QCM_cinema.txt')
    questions_selectionnees = QCM.selectionner_questions(questions, 20)
    questions_melangees = QCM.melanger_reponses(questions_selectionnees)
    
    doc_questions = Document()
    format_paysage(doc_questions)  # Format paysage
    heading = doc_questions.add_heading('Questionnaire à Choix Multiples', 0)
    heading.alignment = 1
    
    # Ajouter une nouvelle section juste après le heading sans passer à une nouvelle page
    section = doc_questions.sections[0]
    set_columns(section)
    
    doc_reponses = Document()
    doc_reponses.add_heading('Réponses du Questionnaire', 0)
    
    reponses_text = []
    
    for index, q in enumerate(questions_melangees):
        if index == 10:
            doc_questions.add_page_break()  # Insérer un saut de page après 10 questions
        p_question = doc_questions.add_paragraph()
        run_question = p_question.add_run(f"{index + 1}. {q['question']}")
        run_question.font.size = Pt(9)
        set_line_spacing(p_question, 1.0, before=2, after=2)
        
        for k, choice in enumerate(q['reponses'], 1):
            lettre = chr(96 + k)
            p = doc_questions.add_paragraph(style='List Bullet')
            run = p.add_run(f"{lettre}. {choice}")
            run.font.size = Pt(9)
            set_line_spacing(p, 1.0, before=1, after=1)
        
        doc_questions.add_paragraph()  # Ajout d'un saut de ligne après chaque question
        
        if q['bonne_reponse'] in q['reponses']:
            reponses_text.append(chr(96 + q['reponses'].index(q['bonne_reponse']) + 1))
    
    # Ajouter un tableau réduit en bas à droite
    table = doc_questions.add_table(rows=2, cols=20)
    table.allow_autofit = True
    table.autofit = False
    table.alignment = 2  # Aligner à droite
    
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(0.25)
            cell.height = Inches(0.15)  # Taille encore plus réduite des cellules
            cell.paragraphs[0].paragraph_format.alignment = 1  # Centrer le texte
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Bordure simple
                border.set(qn('w:sz'), '4')  # Taille fine
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                borders.append(border)
            tcPr.append(borders)
    
    # Remplir la première ligne avec les numéros de 1 à 20
    for i in range(20):
        table.cell(0, i).text = str(i + 1)
        if table.cell(0, i).paragraphs[0].runs:
            table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(9)
        table.cell(0, i).paragraphs[0].paragraph_format.alignment = 1  # Centrer le texte
    
    doc_questions.add_paragraph()  # Espacement minimal pour l'alignement
    doc_reponses.add_paragraph("Réponses :")
    doc_reponses.add_paragraph(", ".join(reponses_text))
    
    doc_questions.save('QCM_Questions.docx')
    doc_reponses.save('QCM_Reponses.docx')

create_qcm()
print("Fichiers créés avec succès !")